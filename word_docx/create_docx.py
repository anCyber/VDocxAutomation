import docx
from docx import Document

from docx.enum.style import WD_STYLE_TYPE           # for custom style configuration
from docx.shared import Pt                          # for setting font sizes
from docx.shared import Mm, Cm                      # for changing paper size and modifying page margins

if (__name__ == "__main__"):
    from functions.functions_table import (
        fill_in_first_row, fill_in_second_row, fill_in_third_row, 
        fill_in_fourth_row, fill_in_fifth_row, fill_in_sixth_row,
        fill_in_seventh_row, fill_in_eigth_row, fill_in_nineth_row
    )
    from functions.functions_text import (
        write_paragraph_organisation_name,
        write_paragraph_recipient_person,
        write_bold_notification_text
    )
else:
    from .functions.functions_table import (
        fill_in_first_row, fill_in_second_row, fill_in_third_row, 
        fill_in_fourth_row, fill_in_fifth_row, fill_in_sixth_row,
        fill_in_seventh_row, fill_in_eigth_row, fill_in_nineth_row
    )
    from .functions.functions_text import (
        write_paragraph_organisation_name,
        write_paragraph_recipient_person,
        write_bold_notification_text
    )


def modify_pages(document: docx.Document) -> None:
    ''' Changes paper size to A4 and modifies margins of all pages in the given file <type: Document>. Original margins: [1.25Cm↑, 1Cm→, 2Cm↓, 2Cm←]. '''
    sections = document.sections
    for section in sections:
        # changing paper size to A4:
        # # Source - https://stackoverflow.com/a/43724253
        # Posted by mkrieger1
        # Retrieved 2026-06-23, License - CC BY-SA 3.0
        section.page_height = Mm(297)
        section.page_width = Mm(210)

        # modifying margins:
        # Source - https://stackoverflow.com/a/32916429
        # Posted by XAnguera
        # Retrieved 2026-06-23, License - CC BY-SA 3.0
        section.top_margin = Cm(1.25)
        section.right_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        
        # actually wrote by myself 
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)



def configure_paragraph_styles(document: docx.Document) -> None:
    ''' Configures the styles later applied to text paragraphs in the given file <type: Document>. Styles: ["TNR12PtStyles", "TNR14PtStyles"] '''
    # Source - https://stackoverflow.com/a/29421050
    # Posted by Alex Nies, modified by community. See post 'Timeline' for change history
    # Retrieved 2026-06-23, License - CC BY-SA 3.0
    for text_size_Pt in [12, 14]:
        font_styles = document.styles
        if (f"TNR{text_size_Pt}PtStyle" in font_styles):
            continue
        else:
            obj_charstyle = font_styles.add_style(f"TNR{text_size_Pt}PtStyle", WD_STYLE_TYPE.CHARACTER)
            obj_font = obj_charstyle.font
            obj_font.size = Pt(text_size_Pt)
            obj_font.name = "Times New Roman"



def create_main_table(document: docx.Document, first_name: str, surname: str, scheduled_date: str, days_amount: int, days_amount_word_form_calendar: str, days_amount_word_form_day: str) -> None:
    ''' Creates, resizes and fills in the table with main info '''
    table = document.add_table(rows=9, cols=6)
    fill_in_first_row(table)
    fill_in_second_row(table, first_name, surname)
    fill_in_third_row(table, scheduled_date, days_amount, days_amount_word_form_calendar, days_amount_word_form_day)
    fill_in_fourth_row(table)
    fill_in_fifth_row(table)
    fill_in_sixth_row(table)
    fill_in_seventh_row(table)
    fill_in_eigth_row(table)
    fill_in_nineth_row(table)



def create_single_page(output_file: docx.Document, should_add_page_break: bool, employee) -> None:
    ''' Creates a single page in the given file <type: Document>. If single page files - do not add page-break.  '''
    write_paragraph_organisation_name(output_file)
    write_paragraph_recipient_person(
        document = output_file,
        position_dative_case = employee.position_dative_case,
        workplace_unit_genitive_case = employee.workplace_unit_genitive_case,
        surname_dative_case = employee.surname_dative_case,
        first_name_initial = employee.first_name_initial,
        patronymic_initial = employee.patronymic_initial
    )
    write_bold_notification_text(output_file)
    create_main_table(
        document = output_file,
        first_name = employee.first_name,
        surname = employee.surname,
        scheduled_date = employee.scheduled_date,
        days_amount = employee.days_amount,
        days_amount_word_form_calendar = employee.days_amount_word_form_calendar,
        days_amount_word_form_day = employee.days_amount_word_form_day
    )
    if (should_add_page_break):
        output_file.add_page_break()

    

def create_many_files(PATH_TO_OUTPUT: str, FILES_AMOUNT: int, list_of_employee_instances: list) -> None:
    ''' Creates N .docx files each containing a single filled page '''
    for i in range (FILES_AMOUNT):
        OutputFile: Document = Document()
        configure_paragraph_styles(OutputFile)
        create_single_page(OutputFile, False, list_of_employee_instances[i])
        modify_pages(OutputFile)
        OutputFile.save(f"{PATH_TO_OUTPUT.removesuffix(".docx")}_{i + 1}.docx")


def create_many_pages(PATH_TO_OUTPUT: str, PAGES_AMOUNT: int, list_of_employee_instances: list) -> None:
    ''' Creates a single .docx file containing N filled pages '''
    OutputFile: Document = Document()
    configure_paragraph_styles(OutputFile)
    for i in range (PAGES_AMOUNT - 1):
        create_single_page(OutputFile, True, list_of_employee_instances[i])
    create_single_page(OutputFile, False, list_of_employee_instances[-1])

    modify_pages(OutputFile)
    OutputFile.save(PATH_TO_OUTPUT)

    



def create_word_document(should_create_single_file: bool, PATH_TO_OUTPUT: str, EMPLOYEES_AMOUNT: int, list_of_employee_instances: list) -> None:
    if (should_create_single_file):
        create_many_pages(PATH_TO_OUTPUT, EMPLOYEES_AMOUNT, list_of_employee_instances)
    else:
        create_many_files(PATH_TO_OUTPUT, EMPLOYEES_AMOUNT, list_of_employee_instances)



if (__name__ == "__main__"):
    create_word_document(True, r"Output\DocxFile.docx", 10)